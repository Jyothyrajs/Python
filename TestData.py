#!/usr/bin/env python3
# -*- coding: utf-8 -*-


import xlrd
from docx import Document
from docx.shared import Inches

class Record :
    def __init__(self, n, i):
        self.name = n
        self.income = i
    def WriteToDoc(self):
        print(self.name, self.income)
        doc = Document()
        doc.add_heading('Income Certificate')
        data = "This is to certify Mr/Ms/Mrs. " + self.name + " has income " + str(self.income)
        doc.add_paragraph(data)
        fileName = self.name + ".docx";
        doc.save(fileName)

def ReadFromExcel():
    wb = xlrd.open_workbook("test.xlsx")
    sheet = wb.sheet_by_index(0)
    for i in range(2, sheet.nrows):
        name = sheet.cell(i,1).value
        income = sheet.cell(i,2).value
        r = Record(name, income)
        r.WriteToDoc()

# Main
ReadFromExcel();
